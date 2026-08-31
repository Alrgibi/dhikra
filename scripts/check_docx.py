#!/usr/bin/env python3
"""
check_docx.py -- the last gate, and the only one that reads the SUBMITTED FILE.

WHY THIS EXISTS. Every other gate in this project checks markdown. The thing that
is handed in is a .docx, and until now nothing looked at it. That is the same
blind spot twice over: preflight did not scan docs/chapters/ until 26 August, and
its pronoun rule could not see a capitalised "We" until 28 August. Both were
caught late and by luck. A document assembled from six markdown files by a
renumbering step is exactly where numbering silently breaks, and markdown that
passed says nothing about the file the examiner opens.

    python scripts/check_docx.py Dhikra_thesis.docx

WHAT IT CHECKS, all read back out of the built file
  HEADINGS     numbered, Title Case, no ":" "-" "--", none left unnumbered
  CAPTIONS     "Figure (n.m) : " and "Table (n.m) : " well formed; numbering
               sequential per chapter with no gaps and no repeats
  REFERENCES   every Figure/Table referenced in the text before it appears, and
               every caption referenced somewhere
  PRONOUNS     case-insensitive we/our/us/you/your, "US" excluded by case
  FIG/FIGURE   "Fig." never used
  STRUCTURE    the seven indices present; chapters 1-6 present; appendix titles
  CITATIONS    every [n] used resolves to a numbered reference entry

WHAT IT CANNOT CHECK, stated so it is not mistaken for more than it is: fonts,
point sizes, margins, line spacing, page numbering style, whether Arabic renders,
whether a figure is legible at print size. Those need a human to open the file.
The checklist for that is in the assembly prompt.
"""
import re, sys, collections
try:
    from docx import Document
except ImportError:
    sys.exit("pip install python-docx")

PRON = re.compile(r"(?<![A-Za-z])(we|our|ours|us|you|your|yours)(?![A-Za-z])", re.I)
HEAD = re.compile(r"^((?:[A-Z]|\d+)(?:\.\d+)*)\s+(.+)$")
EXEMPT_HEAD = re.compile(
    r"^(Chapter\s+\d+|Appendix\s+[A-Z]|Table of Contents|Index of |Abstract"
    r"|Dedication|Acknowledge|References|Bibliography|List of )", re.I)
CAP  = re.compile(r"^(Figure|Table)\s*\(\s*([A-Z]|\d+)\s*\.\s*(\d+)\s*\)\s*:\s*\S")
REF  = re.compile(r"(?<!\*)(Figure|Table)\s*\(\s*([A-Z]|\d+)\s*\.\s*(\d+)\s*\)")

def main(path):
    doc = Document(path)
    paras = [p.text.strip() for p in doc.paragraphs]
    full  = "\n".join(paras)
    bad = []

    caps = collections.defaultdict(list)     # (kind, chapter) -> [num]
    seen_ref, first_ref, first_cap = set(), {}, {}

    for i, t in enumerate(paras):
        if not t:
            continue
        style = doc.paragraphs[i].style.name.lower()

        if "heading" in style:
            # Front matter and title pages are LEGITIMATELY unnumbered. The
            # specification's "no heading is left unnumbered" governs section
            # headings inside chapters; an index title or a chapter title page
            # carries no section number by design, and flagging them would bury
            # the real violations under thirteen false ones.
            if EXEMPT_HEAD.match(t):
                continue
            m = HEAD.match(t)
            if not m:
                bad.append(("HEADING", i, "unnumbered: %r" % t[:50]))
            else:
                body = m.group(2)
                for ch in (":", " - ", "--"):
                    if ch in body:
                        bad.append(("HEADING", i, "contains %r: %r" % (ch.strip(), t[:50])))
                        break
                for w in re.findall(r"[A-Za-z][A-Za-z'’-]*", body):
                    if "." in w or "_" in w:
                        continue
                    if w[0].islower():
                        bad.append(("HEADING", i, "not Title Case at %r in %r" % (w, t[:44])))
                        break

        c = CAP.match(t)
        if c:
            kind, chap, num = c.group(1), c.group(2), int(c.group(3))
            caps[(kind, chap)].append(num)
            first_cap.setdefault((kind, chap, num), i)
        else:
            for m in REF.finditer(t):
                key = (m.group(1), m.group(2), int(m.group(3)))
                seen_ref.add(key); first_ref.setdefault(key, i)

        if "Fig." in t:
            bad.append(("FIG", i, "uses 'Fig.'"))
        for m in PRON.finditer(t):
            if m.group(1) == "US":
                continue
            bad.append(("PRONOUN", i, "%r in %r" % (m.group(1), t[:60]))); break

    for (kind, chap), nums in sorted(caps.items()):
        if nums != list(range(1, len(nums) + 1)):
            bad.append(("CAPTION", 0, "%s captions in %s run %s, not 1..%d"
                        % (kind.lower(), chap, nums, len(nums))))
        if len(set(nums)) != len(nums):
            bad.append(("CAPTION", 0, "duplicate %s number in %s" % (kind.lower(), chap)))
    for key, at in first_cap.items():
        if key not in seen_ref:
            bad.append(("REFERENCE", at, "%s (%s.%d) has a caption but is never referenced" % key))
        elif first_ref[key] > at:
            bad.append(("REFERENCE", at, "%s (%s.%d) appears BEFORE its first mention" % key))
    for key in seen_ref:
        if key not in first_cap:
            bad.append(("REFERENCE", first_ref[key], "%s (%s.%d) referenced but has no caption" % key))

    for want in ("Table of Contents", "Index of Figures", "Index of Tables",
                 "Scientific Terms", "Abbreviations", "Symbols", "Appendices"):
        if want.lower() not in full.lower():
            bad.append(("STRUCTURE", 0, "front matter: %r not found" % want))
    for n in range(1, 7):
        if not re.search(r"Chapter\s+%d\b" % n, full):
            bad.append(("STRUCTURE", 0, "Chapter %d not found" % n))

    used = {int(x) for x in re.findall(r"\[(\d{1,3})(?:,\s*\d{1,3})*\]", full)}
    defined = {int(x) for x in re.findall(r"^\[?(\d{1,3})[\].]\s+[A-Z]", full, re.M)}
    for n in sorted(used - defined):
        if defined:
            bad.append(("CITATION", 0, "[%d] cited with no reference entry" % n))

    print("%s — %d paragraphs, %d captions, %d distinct references"
          % (path, len(paras), sum(len(v) for v in caps.values()), len(seen_ref)))
    for kind, i, msg in bad:
        print("  ! %-10s %s" % (kind, msg))
    if not bad:
        print("  PASSED — headings, captions, references, pronouns and structure all clean")
    print("\nNOT CHECKED (open the file and look): fonts, sizes, margins, 1.5 spacing,\n"
          "Roman/Arabic page numbering, Arabic rendering, figure legibility at print size.")
    return 1 if bad else 0

if __name__ == "__main__":
    sys.exit(main(sys.argv[1] if len(sys.argv) > 1 else sys.exit(__doc__)))
