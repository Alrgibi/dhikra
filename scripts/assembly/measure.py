#!/usr/bin/env python3
"""measure.py <docx> [pattern ...] -- count regex classes over every paragraph AND table cell of a built docx.
   With no patterns: print the standard class table. Reports counts split paragraphs/cells."""
import sys, re, json
from docx import Document

def units(path):
    doc = Document(path)
    out = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if t.strip(): out.append(("p", i, t, (p.style.name if p.style is not None else "")))
    ti = 0
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    if p.text.strip(): out.append(("c", ti, p.text, ""))
        ti += 1
    return out

CLASSES = {
    "dates (d Month yyyy / Month yyyy / yyyy-mm-dd / clock)": r"\b\d{1,2}\s+(January|February|March|April|May|June|July|August|September|October|November|December)\s+20\d\d\b|\b(January|February|March|April|May|June|July|August|September|October|November|December)\s+20\d\d\b|\b20\d\d-\d\d-\d\d\b|\b\d{1,2}:\d\d\b",
    "heading punctuation (headings only)": None,   # computed separately
    "external-corpus variants": r"never seen|untouched|never read|withheld from every|excluded from every development decision|held out from the entire",
    "scaffolding": r"local to this appendix|References Cited In This|scaffolding",
    "[DRAFT]": r"\[DRAFT",
    "author-year brackets": r"\[(Sedighi|Dajani|El-Metwally|Hanley)[^\]]*\d{4}\]",
    "64-feature": r"64-feature",
    "117": r"\b117\b",
    "5.1.3": r"\b5\.1\.3\b",
    "[UNVERIFIED]": r"\[UNVERIFIED\]",
    "33.3%": r"33\.3\s?%|0\.3333",
    "77.8%": r"77\.8\s?%|0\.7778",
    "0.547": r"0\.547\b|0\.5474",
    "0.6291": r"0\.6291|0\.629\b",
    "INADEQUATE": r"INADEQUATE",
    "rather than assumed": r"rather than assumed",
    "is reported rather than": r"is reported rather than",
    "was measured rather than": r"was measured rather than",
    "measured rather than assumed": r"measured rather than assumed",
    "pre-registered": r"pre-registered|preregistered",
}

def heading_punct(path):
    doc = Document(path)
    bad = []
    for p in doc.paragraphs:
        st = p.style.name if p.style is not None else ""
        if "Heading" in st and re.match(r"^(\d+|[A-Z])(\.\d+)+\s", p.text):
            if re.search(r"[,:—–]| - |--|\s-\s", p.text) or re.search(r"(?<=[A-Za-z])-(?=[A-Za-z])", p.text):
                bad.append(p.text)
    return bad

if __name__ == "__main__":
    path = sys.argv[1]
    U = units(path)
    if len(sys.argv) > 2:
        for pat in sys.argv[2:]:
            rx = re.compile(pat)
            for kind, i, t, st in U:
                for m in rx.finditer(t):
                    a, b = max(0, m.start() - 70), min(len(t), m.end() + 50)
                    print(f"[{kind}{i}] …{t[a:b]}…")
        sys.exit(0)
    print(f"{path}: {sum(1 for u in U if u[0]=='p')} paragraphs, {sum(1 for u in U if u[0]=='c')} cell paragraphs")
    for name, pat in CLASSES.items():
        if pat is None:
            hp = heading_punct(path); print(f"{name:52s} {len(hp):4d}"); continue
        rx = re.compile(pat)
        cp = sum(len(rx.findall(t)) for k, i, t, s in U if k == "p")
        cc = sum(len(rx.findall(t)) for k, i, t, s in U if k == "c")
        print(f"{name:52s} {cp+cc:4d}  (paragraphs {cp}, cells {cc})")
